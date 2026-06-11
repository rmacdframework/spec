"""Regenerate the Word copy of the spec from the authoritative Markdown.

The Markdown file (``RMACD_Framework_v1.4.md``) is the document of record;
the ``.docx`` is a distribution convenience. This converter handles the
constructs the spec actually uses:

- ATX headings ``#``..``####`` (bold markers inside headings are stripped)
- paragraphs with inline ``**bold**``, ``*italic*``, ``code`` and links
  (link text is kept, the URL is dropped for footnote-free print copy)
- ``-`` bullet lists and ``1.`` numbered lists
- pipe tables (header row bold, Table Grid style)
- fenced code blocks (monospace, shaded-free, preserved verbatim)
- images ``![alt](relative.png)`` (6.5in wide)
- horizontal rules (skipped — Word section spacing covers it)

Usage::

    python convert_md_to_docx.py RMACD_Framework_v1.4.md
    # writes RMACD_Framework_v1.4.docx next to the input
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

_INLINE = re.compile(
    r"(\*\*.+?\*\*"      # bold
    r"|\*[^*\n]+?\*"     # italic
    r"|`[^`\n]+?`"       # inline code
    r"|\[[^\]]+\]\([^)]+\))"  # link
)
_IMG = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")
_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET = re.compile(r"^[-*]\s+(.*)$")
_NUMBERED = re.compile(r"^\d+[.)]\s+(.*)$")


def add_runs(paragraph, text: str) -> None:
    """Split inline markdown into styled runs."""
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("*"):
            paragraph.add_run(token[1:-1]).italic = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
        else:  # [text](url) — keep the text only
            paragraph.add_run(token[1:token.index("]")])
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(8)
    for i, line in enumerate(lines):
        run = p.add_run(line + ("\n" if i < len(lines) - 1 else ""))
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)


def add_table(doc: Document, rows: list[str]) -> None:
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
            continue  # separator row
        parsed.append(cells)
    if not parsed:
        return
    n_cols = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=n_cols)
    table.style = "Table Grid"
    for r, row in enumerate(parsed):
        for c in range(n_cols):
            cell = table.cell(r, c)
            cell.paragraphs[0].text = ""
            add_runs(cell.paragraphs[0], row[c] if c < len(row) else "")
            if r == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9.5)
    doc.add_paragraph()


def convert(md_path: Path, out_path: Path | None = None) -> Path:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    for section in doc.sections:  # a bit more room for tables
        section.left_margin = section.right_margin = Inches(0.9)

    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            block: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i])
                i += 1
            add_code_block(doc, block)
            i += 1
            continue

        if line.strip().startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(lines[i])
                i += 1
            add_table(doc, rows)
            continue

        if m := _HEADING.match(line):
            level = min(len(m.group(1)), 4)
            text = m.group(2).strip().replace("**", "")
            doc.add_heading(text, level=level)
            i += 1
            continue

        if m := _IMG.match(line.strip()):
            img = (md_path.parent / m.group(2)).resolve()
            if img.exists():
                doc.add_picture(str(img), width=Inches(6.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = doc.add_paragraph()
                add_runs(p, f"[image: {m.group(1) or m.group(2)}]")
            i += 1
            continue

        if re.fullmatch(r"\s*([-*_]\s*){3,}", line):
            i += 1  # horizontal rule
            continue

        if m := _BULLET.match(line.strip()):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, m.group(1))
            i += 1
            continue

        if m := _NUMBERED.match(line.strip()):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, m.group(1))
            i += 1
            continue

        if line.strip():
            p = doc.add_paragraph()
            add_runs(p, line.strip())
        i += 1

    out = out_path or md_path.with_suffix(".docx")
    doc.save(out)
    return out


def main() -> None:
    if len(sys.argv) < 2:
        print(f"usage: {sys.argv[0]} <file.md>", file=sys.stderr)
        sys.exit(2)
    out = convert(Path(sys.argv[1]))
    print(f"wrote {out}")


if __name__ == "__main__":
    main()
